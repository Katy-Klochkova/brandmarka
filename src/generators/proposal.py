"""Генерация коммерческого предложения в формате DOCX в стиле Anthropic."""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.brand import load as load_brand
from core.config import OUTPUT_DIR
from core.exporter import office_to_pdf


_DEFAULT_PRIMARY = "#1F1F1E"
_DEFAULT_SECONDARY = "#5A4BFF"
_DEFAULT_ACCENT = "#FF5B24"
_DEFAULT_TEXT = "#1F1F1E"


def _hex_to_rgb(hex_color: str) -> RGBColor:
    """Преобразовать HEX в RGBColor."""
    hex_color = hex_color.lstrip("#")
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def _brand_colors(brand: dict) -> dict:
    """Цвета бренда с антропик-фолбэками."""
    colors = brand.get("colors", {})
    return {
        "primary": _hex_to_rgb(colors.get("primary", _DEFAULT_PRIMARY)),
        "secondary": _hex_to_rgb(colors.get("secondary", _DEFAULT_SECONDARY)),
        "accent": _hex_to_rgb(colors.get("accent", _DEFAULT_ACCENT)),
        "text": _hex_to_rgb(colors.get("text", _DEFAULT_TEXT)),
    }


def _set_cell_shading(cell, fill: str) -> None:
    """Заливка ячейки таблицы."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def _set_cell_border(cell, color: str) -> None:
    """Установить границу ячейки."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = 'w:{}'.format(edge)
        element = OxmlElement(tag)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '6')
        element.set(qn('w:color'), color)
        tcBorders.append(element)
    tcPr.append(tcBorders)


def generate(input_data: dict, output_dir: Path = OUTPUT_DIR, output_format: str = "docx") -> Path:
    """Сгенерировать коммерческое предложение и вернуть путь к файлу."""
    brand = load_brand()
    output_dir.mkdir(parents=True, exist_ok=True)

    client_name = input_data.get("client_name", "")
    services = input_data.get("services", [])
    terms = input_data.get("terms", "")
    payment = input_data.get("payment", "")
    company = brand.get("company", "")
    slogan = brand.get("slogan", "")
    contacts = brand.get("contacts", {})

    colors = _brand_colors(brand)

    doc = Document()

    # Узкие поля
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Заголовок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Коммерческое предложение")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = colors["primary"]

    # Мета
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"{company}\n")
    run.font.size = Pt(12)
    run.font.color.rgb = colors["text"]
    if slogan:
        run = meta.add_run(slogan)
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = colors["secondary"]

    # Кому
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    run = p.add_run(f"Для: {client_name}")
    run.font.bold = True
    run.font.color.rgb = colors["primary"]
    run.font.size = Pt(12)

    # Введение
    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(8)
    intro.paragraph_format.space_after = Pt(12)
    run = intro.add_run(
        f"{company} предлагает выполнить работы для {client_name} в соответствии с нижеуказанным перечнем услуг и условиями."
    )
    run.font.size = Pt(11)
    run.font.color.rgb = colors["text"]

    # Таблица услуг
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(0.6)
    table.columns[1].width = Inches(4.5)
    table.columns[2].width = Inches(1.5)

    hdr_cells = table.rows[0].cells
    headers = ["№", "Услуга", "Стоимость, ₽"]
    for i, text in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
        _set_cell_shading(hdr_cells[i], "5A4BFF")
        _set_cell_border(hdr_cells[i], "5A4BFF")

    total = 0
    for i, service in enumerate(services, start=1):
        row_cells = table.add_row().cells
        price = service.get("price", 0)
        total += price
        values = [str(i), service.get("name", ""), f"{price:,.0f}".replace(",", " ")]
        for j, text in enumerate(values):
            row_cells[j].text = ""
            p = row_cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.color.rgb = colors["text"]
            _set_cell_border(row_cells[j], "E8E8E6")

    # Итого
    row_cells = table.add_row().cells
    row_cells[0].text = ""
    row_cells[1].text = ""
    row_cells[2].text = ""
    p = row_cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Итого")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = colors["primary"]
    p = row_cells[2].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{total:,.0f}".replace(",", " "))
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = colors["primary"]
    _set_cell_shading(row_cells[1], "EDE9FF")
    _set_cell_shading(row_cells[2], "EDE9FF")
    _set_cell_border(row_cells[1], "5A4BFF")
    _set_cell_border(row_cells[2], "5A4BFF")

    # Сроки и оплата
    doc.add_paragraph()
    if terms:
        p = doc.add_paragraph()
        run = p.add_run("Сроки выполнения: ")
        run.font.bold = True
        run.font.color.rgb = colors["primary"]
        run.font.size = Pt(11)
        run = p.add_run(terms)
        run.font.size = Pt(11)
        run.font.color.rgb = colors["text"]

    if payment:
        p = doc.add_paragraph()
        run = p.add_run("Условия оплаты: ")
        run.font.bold = True
        run.font.color.rgb = colors["primary"]
        run.font.size = Pt(11)
        run = p.add_run(payment)
        run.font.size = Pt(11)
        run.font.color.rgb = colors["text"]

    # Контакты
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Контакты для связи")
    run.font.bold = True
    run.font.color.rgb = colors["secondary"]
    run.font.size = Pt(12)
    contact_lines = [
        contacts.get("phone", ""),
        contacts.get("email", ""),
        contacts.get("site", ""),
    ]
    for line in [l for l in contact_lines if l]:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = colors["text"]

    safe_name = client_name.replace(" ", "_").replace(".", "")
    docx_path = output_dir / f"proposal_{safe_name}.docx"
    doc.save(docx_path)

    if output_format.lower() == "pdf":
        return office_to_pdf(docx_path, output_dir)
    return docx_path


if __name__ == "__main__":
    sample = {
        "client_name": "ООО РоторПром",
        "services": [
            {"name": "Поставка системы вибродиагностики", "price": 1250000},
            {"name": "Монтаж и пусконаладка", "price": 280000},
            {"name": "Обучение персонала", "price": 120000},
        ],
        "terms": "30 рабочих дней с момента подписания договора",
        "payment": "50% предоплата, 50% после ввода в эксплуатацию",
    }
    path = generate(sample)
    print(f"КП сохранено: {path}")
